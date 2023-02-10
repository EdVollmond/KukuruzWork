import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Mm
import json

def table_writer(ITEMS0,ITEMS1,docname):

    with open('JSON/turnprotectiontime.json',  'r', encoding='utf-8') as fr:
        turnprotectiontime = json.load(fr)

    doc = docx.Document('acts/' + docname)

    table0 = doc.tables[0]
    table1 = doc.tables[1]

    sizcount=0

    if turnprotectiontime == True:
        for partynum in range (0,len(ITEMS0)):
            for rownum in range (0,len(ITEMS0[partynum])):
                if ITEMS0[partynum][rownum][-1] != "> 25":
                    ITEMS0[partynum][rownum].append("> 25")

    for partynum in range (0,len(ITEMS0)):
        for rownum in range (0,len(ITEMS0[partynum])):
            row = table1.add_row().cells
            for cellnum in range (0,len(ITEMS0[partynum][rownum])):
                line = rownum+sizcount
                cell = str(ITEMS0[partynum][rownum][cellnum])

                cols = table1.rows[line + 2].cells
                text = cols[cellnum].paragraphs[0]
                text.add_run(cell).font.size = Pt(11)
                text.first_line_indent = Mm(0)
                text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                

        sizcount = sizcount + len(ITEMS0[partynum])

    for rownum in range (0,len(ITEMS1)):
        row = table0.add_row().cells
        for cellnum in range (0,len(ITEMS1[rownum])):
            table0.cell(rownum+1, cellnum).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if cellnum == 1:
                cell = str(ITEMS1[rownum][cellnum])
                cell = cell.replace("противогазы","Противогазы")
                cell = cell.replace("доп","Доп")
                cols = table0.rows[rownum + 1].cells
                text = cols[cellnum].paragraphs[0]
                text.add_run(cell).font.size = Pt(14)
                text.first_line_indent = Mm(0)
                text.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell = str(ITEMS1[rownum][cellnum])
                cols = table0.rows[rownum + 1].cells
                text = cols[cellnum].paragraphs[0]
                text.add_run(cell).font.size = Pt(14)
                text.first_line_indent = Mm(0)
                text.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.save('acts/' + docname)
    print("Табличные данные введены")
